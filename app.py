import streamlit as st
# ============== 修复包版本缺失的补丁 ==============
import sys
import importlib.metadata
# 获取原来的版本读取函数
_orig_version = importlib.metadata.version
# 自定义一个新函数：当程序问 streamlit 的版本时，直接给它一个假版本号
def _patched_version(package_name):
    if package_name == 'streamlit':
        return '1.71.0'
    return _orig_version(package_name)
# 把原函数替换掉
importlib.metadata.version = _patched_version
# =================================================

import streamlit as st
# (接着就是你下面原本的所有代码)
import sqlite3
import bcrypt
import os
import shutil
from datetime import datetime, date, time
import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
from openpyxl import Workbook
from pypinyin import pinyin

from PIL import Image
import io

def compress_image(file_bytes, max_size_mb=2):
    """压缩图片至 max_size_mb 以内，返回压缩后的字节数据"""
    try:
        img = Image.open(io.BytesIO(file_bytes))
        # 将图片转为 JPEG 格式（如果原图是 PNG 也可能被压缩）
        quality = 85
        output = io.BytesIO()
        img.save(output, format='JPEG', quality=quality, optimize=True)
        while output.tell() > max_size_mb * 1024 * 1024 and quality > 20:
            quality -= 5
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=quality, optimize=True)
        return output.getvalue()
    except Exception as e:
        print(f"图片压缩失败: {e}")
        return file_bytes  # 压缩失败返回原数据
# ================= 页面配置 =================
st.set_page_config(page_title="知微 - 你的私人AI助理", layout="wide")

# ================= 存储路径配置 =================
BASE_STORAGE = "D:/知微AI"
try:
    os.makedirs(BASE_STORAGE, exist_ok=True)
except:
    BASE_STORAGE = os.getcwd() + "/知微AI"
    os.makedirs(BASE_STORAGE, exist_ok=True)

def get_user_storage(user_id):
    user_dir = os.path.join(BASE_STORAGE, f"user_{user_id}")
    sub_dirs = ["notes_files", "contacts_avatars", "projects_files", "atom_notes_files"]
    for sub in sub_dirs:
        os.makedirs(os.path.join(user_dir, sub), exist_ok=True)
    return user_dir

# ================= 数据库初始化（带自动升级） =================
DB_PATH = os.path.join(BASE_STORAGE, "zhiwei.db")

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    # 1. 创建老表结构（如果不存在）
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE, password_hash TEXT, identity_code TEXT UNIQUE, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS notes
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, content TEXT, tags TEXT, file_paths TEXT, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS contacts
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, contact_code TEXT, name TEXT, company TEXT, department TEXT, position TEXT, phone TEXT, avatar_path TEXT, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS projects
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, contact_id INTEGER, project_name TEXT, file_paths TEXT, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS project_folders
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, contact_id INTEGER, folder_name TEXT, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS atom_notes
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, title TEXT, content TEXT, file_paths TEXT, code_content TEXT, created_at TEXT)''')
    
    # 2. 自动检测并升级数据库
    c.execute("PRAGMA table_info(notes)")
    columns = [col[1] for col in c.fetchall()]
    if 'title' not in columns:
        c.execute("ALTER TABLE notes ADD COLUMN title TEXT DEFAULT '无标题'")
    if 'reminder' not in columns:
        c.execute("ALTER TABLE notes ADD COLUMN reminder TEXT DEFAULT ''")
    
    c.execute("PRAGMA table_info(atom_notes)")
    atom_columns = [col[1] for col in c.fetchall()]
    if 'code_content' not in atom_columns:
        c.execute("ALTER TABLE atom_notes ADD COLUMN code_content TEXT DEFAULT ''")
        
    conn.commit()
    conn.close()

init_db()

# ================= 辅助函数 =================
def hash_password(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def check_password(password, hashed):
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def get_db_connection():
    return sqlite3.connect(DB_PATH)

def get_identity_prefix(username):
    if not username:
        return "def"
    first_char = username[0]
    if '\u4e00' <= first_char <= '\u9fff': # 是中文
        try:
            pinyin_list = pinyin(first_char)
            return pinyin_list[0][0][:3].lower()
        except:
            return username[:3].lower()
    else:
        return username[:3].upper()

# ================= 登录/注册逻辑 =================
def login_page():
    st.title("知微 · 登录")
    with st.form("login_form"):
        username = st.text_input("用户名")
        password = st.text_input("密码", type="password")
        submit = st.form_submit_button("登录")
        
        if submit:
            conn = get_db_connection()
            c = conn.cursor()
            c.execute("SELECT id, password_hash FROM users WHERE username = ?", (username.strip(),))
            user = c.fetchone()
            conn.close()
            if user and check_password(password.strip(), user[1]):
                st.session_state["user"] = {"id": user[0], "username": username.strip()}
                st.success("登录成功！")
                st.rerun()
            else:
                st.error("用户名或密码错误，请尝试重新注册（若在云端运行，数据每次重启会重置）。")

    st.divider()
    st.write("还没有账号？")
    if st.button("去注册"):
        st.session_state["page"] = "register"
        st.rerun()

def register_page():
    st.title("知微 · 注册")
    with st.form("register_form"):
        new_username = st.text_input("设置用户名")
        new_password = st.text_input("设置密码", type="password")
        confirm_password = st.text_input("确认密码", type="password")
        submit = st.form_submit_button("注册")
        
        if submit:
            clean_username = new_username.strip()
            clean_password = new_password.strip()
            clean_confirm = confirm_password.strip()

            if not clean_username or not clean_password:
                st.error("用户名和密码不能为空")
            elif clean_password != clean_confirm:
                st.error("两次密码不一致")
            else:
                conn = get_db_connection()
                c = conn.cursor()
                try:
                    prefix = get_identity_prefix(clean_username)
                    c.execute("SELECT MAX(id) FROM users")
                    max_id = c.fetchone()[0]
                    next_id = 1 if max_id is None else max_id + 1
                    identity_code = prefix + str(next_id).zfill(15)
                    
                    hashed = hash_password(clean_password)
                    c.execute("INSERT INTO users (username, password_hash, identity_code, created_at) VALUES (?, ?, ?, ?)",
                              (clean_username, hashed, identity_code, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                    conn.commit()
                    st.success(f"注册成功！您的专属身份码是：**{identity_code}**")
                    st.session_state["page"] = "login"
                    st.rerun()
                except sqlite3.IntegrityError:
                    st.error("用户名已存在，或身份码冲突。")
                finally:
                    conn.close()

    if st.button("返回登录"):
        st.session_state["page"] = "login"
        st.rerun()

# ================= 主应用逻辑 =================
def main_app():
    user = st.session_state["user"]
    user_storage = get_user_storage(user['id'])
    
    st.sidebar.title(f"👋 你好，{user['username']}")
    menu = st.sidebar.radio("功能导航", ["📝 我的笔记", "📅 周报生成", "📊 思维导图", "🔍 联网搜索", "🤖 陪伴对话", "👤 联系人管理", "📒 原子笔记", "⚙️ 个人中心"])
    
    if st.sidebar.button("退出登录"):
        st.session_state.clear()
        st.rerun()

    # ---------- 笔记模块 ----------
    if menu == "📝 我的笔记":
        st.subheader("📝 我的笔记")
        
        search_query = st.text_input("🔍 按标题/标签快速搜索笔记")
        
        with st.expander("✏️ 写新笔记 / 上传附件", expanded=True):
            with st.form("new_note"):
                title = st.text_input("笔记标题（文件命名）")
                content = st.text_area("内容 (支持 Markdown 语法: # 标题, **加粗**, ```代码```)", height=150)
                tags = st.text_input("标签（用逗号分隔）")
                
                col_date, col_time = st.columns(2)
                with col_date:
                    note_date = st.date_input("提醒日期", value=date.today())
                with col_time:
                    note_time = st.time_input("提醒时间", value=time(hour=9, minute=0))
                
                uploaded_files = st.file_uploader("上传附件", type=['docx','xlsx','pdf','pptx','jpg','png'], accept_multiple_files=True)
                submitted = st.form_submit_button("保存笔记")
                
                if submitted and title:
                    file_paths = []
                    if uploaded_files:
                        notes_dir = os.path.join(user_storage, "notes_files")
                        for file in uploaded_files:
                            save_path = os.path.join(notes_dir, file.name)
                            with open(save_path, "wb") as f:
                                f.write(file.getbuffer())
                            file_paths.append(save_path)
                    file_paths_str = ",".join(file_paths)
                    reminder_str = note_date.strftime("%Y-%m-%d") + " " + note_time.strftime("%H:%M")
                    
                    conn = get_db_connection()
                    c = conn.cursor()
                    c.execute("INSERT INTO notes (user_id, title, content, tags, file_paths, reminder, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                              (user['id'], title, content, tags, file_paths_str, reminder_str, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                    conn.commit()
                    conn.close()
                    st.success("✅ 笔记及附件已成功保存！")
                    st.rerun()

        conn = get_db_connection()
        c = conn.cursor()
        if search_query:
            c.execute('''SELECT id, title, content, tags, file_paths, reminder, created_at FROM notes 
                         WHERE user_id = ? AND (title LIKE ? OR tags LIKE ?) ORDER BY created_at DESC''', 
                         (user['id'], f"%{search_query}%", f"%{search_query}%"))
        else:
            c.execute("SELECT id, title, content, tags, file_paths, reminder, created_at FROM notes WHERE user_id = ? ORDER BY created_at DESC", (user['id'],))
        notes = c.fetchall()
        conn.close()

        if not notes:
            st.info("还没有笔记，写一条吧！")
        else:
            for note in notes:
                with st.container(border=True):
                    st.markdown(f"### {note[1]}") # 标题
                    st.write(f"**{note[6]}**") # 时间
                    if note[5]: # 提醒时间
                        st.caption(f"⏰ 提醒：{note[5]}")
                    
                    st.markdown(note[2][:300] + "..." if len(note[2]) > 300 else note[2])
                    
                    if note[3]:
                        st.caption(f"🏷️ {note[3]}")
                    if note[4]:
                        st.write("📎 **附件：**")
                        for path in note[4].split(","):
                            if path:
                                fname = os.path.basename(path)
                                if os.path.exists(path):
                                    with open(path, "rb") as f:
                                        st.download_button(label=f"📥 下载 {fname}", data=f, file_name=fname, key=f"dl_{note[0]}_{fname}")
                    if st.button(f"删除 #{note[0]}", key=f"del_{note[0]}"):
                        conn = get_db_connection()
                        c = conn.cursor()
                        c.execute("DELETE FROM notes WHERE id = ? AND user_id = ?", (note[0], user['id']))
                        conn.commit()
                        conn.close()
                        st.rerun()

    # ---------- 周报生成 ----------
    elif menu == "📅 周报生成":
        st.subheader("📅 本周周报生成")
        if st.button("🔄 一键生成本周周报"):
            conn = get_db_connection()
            c = conn.cursor()
            c.execute('''SELECT title, content, created_at FROM notes 
                         WHERE user_id = ? AND date(created_at) >= date('now', 'weekday 1', '-7 days')
                         ORDER BY created_at DESC''', (user['id'],))
            week_notes = c.fetchall()
            conn.close()
            
            if not week_notes:
                st.warning("本周暂无笔记记录。")
            else:
                report = f"# 知微 · 本周工作周报\n\n**生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n## 本周主要事项\n"
                for idx, note in enumerate(week_notes, 1):
                    report += f"{idx}. **{note[0]}**: {note[1][:50]}... (记录于 {note[2]})\n"
                st.markdown(report)
                
                doc = Document()
                doc.add_heading('知微 · 本周工作周报', 0)
                doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_heading('本周主要事项', level=1)
                for idx, note in enumerate(week_notes, 1):
                    doc.add_paragraph(f"{idx}. {note[0]}: {note[1][:50]}...", style='List Bullet')
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                wb = Workbook()
                ws = wb.active
                ws.title = "本周周报"
                ws.append(["序号", "标题", "内容摘要", "记录时间"])
                for idx, note in enumerate(week_notes, 1):
                    ws.append([idx, note[0], note[1][:50], note[2]])
                excel_io = BytesIO()
                wb.save(excel_io)
                excel_io.seek(0)

                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(label="📄 导出 Word (.docx)", data=doc_io, file_name="知微AI_本周周报.docx")
                with col2:
                    st.download_button(label="📊 导出 Excel (.xlsx)", data=excel_io, file_name="知微AI_本周周报.xlsx")

    # ---------- 思维导图 ----------
    elif menu == "📊 思维导图":
        st.info("💡 功能开发中：请确保本地安装了系统级 Graphviz。")

    # ---------- 联网搜索 ----------
    elif menu == "🔍 联网搜索":
        st.subheader("🔍 联网搜索（Bing）")
        query = st.text_input("请输入要搜索的关键词")
        if st.button("开始搜索") and query:
            with st.spinner(f"正在搜索：{query} ..."):
                try:
                    headers = {'User-Agent': 'Mozilla/5.0'}
                    url = f"https://www.bing.com/search?q={query}"
                    response = requests.get(url, headers=headers, timeout=10)
                    soup = BeautifulSoup(response.text, 'html.parser')
                    results = []
                    for li in soup.select('li.b_algo h2 a'):
                        href = li.get('href')
                        if href and href.startswith('http'):
                            results.append(href)
                    if results:
                        st.success(f"找到 {len(results)} 条结果：")
                        for i, url in enumerate(results[:5], 1):
                            st.write(f"{i}. [{url}]({url})")
                    else:
                        st.warning("未找到相关结果。")
                except Exception as e:
                    st.error(f"搜索失败：{e}")

    # ---------- 陪伴对话 ----------
    elif menu == "🤖 陪伴对话":
        st.subheader("🤖 知微助手（轻量版）")
        if "chat_msg" not in st.session_state:
            st.session_state.chat_msg = []
        user_input = st.text_input("对我说句话吧")
        if st.button("发送", key="chat_send") and user_input:
            if "忙" in user_input or "累" in user_input:
                reply = "辛苦了！建议您梳理一下笔记，偶尔放空也是生产力。"
            elif "思路" in user_input or "计划" in user_input:
                reply = "您可以试试用'笔记'把想法写下来。"
            else:
                reply = "收到！您可以在'笔记'或'原子笔记'中记录详细内容。"
            st.session_state.chat_msg.append(("你", user_input))
            st.session_state.chat_msg.append(("知微", reply))
        for role, text in st.session_state.chat_msg:
            if role == "知微":
                st.info(f"🤖 **知微**：{text}")
            else:
                st.write(f"👤 **你**：{text}")

    # ---------- 联系人管理 ----------
    elif menu == "👤 联系人管理":
        st.subheader("👤 联系人管理")
        
        with st.expander("➕ 添加新联系人", expanded=False):
            with st.form("add_contact"):
                col1, col2 = st.columns(2)
                with col1:
                    c_name = st.text_input("姓名 *")
                    c_company = st.text_input("公司")
                    c_dept = st.text_input("部门")
                    c_pos = st.text_input("职务")
                with col2:
                    c_phone = st.text_input("联系方式")
                    c_code_input = st.text_input("粘贴对方18位身份码")
                c_submit = st.form_submit_button("保存联系人")
                if c_submit and c_name:
                    conn = get_db_connection()
                    c = conn.cursor()
                    c.execute("INSERT INTO contacts (user_id, contact_code, name, company, department, position, phone, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                              (user['id'], c_code_input, c_name.strip(), c_company.strip(), c_dept.strip(), c_pos.strip(), c_phone.strip(), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                    conn.commit()
                    conn.close()
                    st.success("联系人添加成功！")
                    st.rerun()

        search_query = st.text_input("🔍 快速搜索联系人")
        
        conn = get_db_connection()
        c = conn.cursor()
        if search_query:
            c.execute('''SELECT id, name, company, department, position, phone, contact_code, created_at 
                         FROM contacts WHERE user_id = ? AND name LIKE ? ORDER BY created_at DESC''', 
                         (user['id'], f"%{search_query}%"))
        else:
            c.execute("SELECT id, name, company, department, position, phone, contact_code, created_at FROM contacts WHERE user_id = ? ORDER BY created_at DESC", (user['id'],))
        contacts = c.fetchall()
        conn.close()

        if not contacts:
            st.info("暂无联系人。")
        else:
            for contact in contacts:
                cid, name, company, dept, pos, phone, code, c_time = contact
                with st.container(border=True):
                    cols = st.columns([1, 6, 2])
                    with cols[0]:
                        st.image(f"https://ui-avatars.com/api/?name={name}&background=random&color=fff&size=64", use_container_width=True)
                    with cols[1]:
                        st.markdown(f"**{name}**")
                        st.caption(f"🏢 {company} | {dept} | {pos} | 📞 {phone}")
                        st.caption(f"🆔 身份码: {code}")
                        
                        st.write("📂 **历史项目（文件夹）：**")
                        conn = get_db_connection()
                        c = conn.cursor()
                        c.execute("SELECT id, folder_name, created_at FROM project_folders WHERE user_id = ? AND contact_id = ? ORDER BY created_at DESC", (user['id'], cid))
                        folders = c.fetchall()
                        conn.close()
                        
                        if folders:
                            for folder in folders:
                                fid, fname, ftime = folder
                                st.write(f"- 🗂️ **{fname}** ({ftime})")
                                conn = get_db_connection()
                                c = conn.cursor()
                                c.execute("SELECT id, project_name, file_paths, created_at FROM projects WHERE folder_id = ? ORDER BY created_at DESC", (fid,))
                                projects = c.fetchall()
                                conn.close()
                                for proj in projects:
                                    st.write(f"  - 📄 **{proj[1]}** ({proj[3]})")
                                    if proj[2]:
                                        for fpath in proj[2].split(","):
                                            if fpath and os.path.exists(fpath):
                                                fname_file = os.path.basename(fpath)
                                                with open(fpath, "rb") as f:
                                                    st.download_button(label=f"📎 下载 {fname_file}", data=f, file_name=fname_file, key=f"proj_{proj[0]}_{fname_file}")
                        else:
                            st.write("*暂无文件夹/项目*")
                    
                    with cols[2]:
                        with st.popover("➕ 新增项目文件夹"):
                            new_folder_name = st.text_input("文件夹名称", key=f"folder_name_{cid}")
                            if st.button("创建文件夹", key=f"create_folder_{cid}") and new_folder_name:
                                conn = get_db_connection()
                                c = conn.cursor()
                                c.execute("INSERT INTO project_folders (user_id, contact_id, folder_name, created_at) VALUES (?, ?, ?, ?)",
                                          (user['id'], cid, new_folder_name.strip(), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                                conn.commit()
                                conn.close()
                                st.success("文件夹已创建！")
                                st.rerun()
                        
                        with st.popover("➕ 新增项目"):
                            all_folders = []
                            conn = get_db_connection()
                            c = conn.cursor()
                            c.execute("SELECT id, folder_name FROM project_folders WHERE user_id = ? AND contact_id = ?", (user['id'], cid))
                            all_folders = c.fetchall()
                            conn.close()
                            
                            if not all_folders:
                                st.warning("请先创建文件夹！")
                            else:
                                folder_options = {f[1]: f[0] for f in all_folders}
                                selected_folder_name = st.selectbox("选择文件夹", list(folder_options.keys()), key=f"folder_sel_{cid}")
                                proj_name = st.text_input("项目名称", key=f"proj_name_{cid}")
                                proj_files = st.file_uploader("上传项目附件", accept_multiple_files=True, key=f"proj_files_{cid}")
                                if st.button("保存项目", key=f"save_proj_{cid}") and proj_name:
                                    proj_dir = os.path.join(user_storage, "projects_files")
                                    file_paths = []
                                    for f in proj_files:
                                        save_path = os.path.join(proj_dir, f.name)
                                        with open(save_path, "wb") as f_out:
                                            f_out.write(f.getbuffer())
                                        file_paths.append(save_path)
                                    conn = get_db_connection()
                                    c = conn.cursor()
                                    c.execute("INSERT INTO projects (contact_id, folder_id, project_name, file_paths, created_at) VALUES (?, ?, ?, ?, ?)",
                                              (cid, folder_options[selected_folder_name], proj_name.strip(), ",".join(file_paths), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                                    conn.commit()
                                    conn.close()
                                    st.success("项目已添加！")
                                    st.rerun()

    # ---------- 原子笔记（彻底解决 DOM 崩溃问题） ----------
    # ---------- 原子笔记（图文序列 + Drawflow 流程图） ----------
    elif menu == "📒 原子笔记":
        # 初始化流程图状态
        if "flow_json" not in st.session_state:
            st.session_state.flow_json = "{}"

        tab_note, tab_flow = st.tabs(["📝 图文笔记编辑", "📊 白板流程图 (WPS级零代码)"])

        with tab_note:
            st.subheader("📒 原子笔记（图文序列与分屏对比）")
            st.caption("支持 Markdown 图文混排，下方可零代码点击构建流程图。")
            
            if "atom_content_prompt" not in st.session_state:
                st.session_state.atom_content_prompt = ""
            
            with st.expander("✏️ 新建图文笔记 / 制作记录", expanded=True):
                col_btn1, col_btn2, col_btn3, col_btn4, col_btn5 = st.columns(5)
                if col_btn1.button("🔲 待办", key="todo_btn"):
                    st.session_state.atom_content_prompt += "\n- [ ] 待办事项\n- [ ] 任务2"
                if col_btn2.button("**B** 粗体", key="bold_btn"):
                    st.session_state.atom_content_prompt += "**加粗文本**"
                if col_btn3.button("*I* 斜体", key="italic_btn"):
                    st.session_state.atom_content_prompt += "*斜体文本*"
                if col_btn4.button("S 删线", key="strike_btn"):
                    st.session_state.atom_content_prompt += "~~删除文本~~"
                if col_btn5.button("📝 标题", key="header_btn"):
                    st.session_state.atom_content_prompt += "\n### 新章节"
                    
                with st.form("atom_note"):
                    a_title = st.text_input("笔记标题")
                    c1_img, c2_img = st.columns(2)
                    with c1_img:
                        upload_img = st.file_uploader("📁 上传图片", type=['png','jpg','jpeg','gif'], key="atom_img_upload")
                    with c2_img:
                        camera_img = st.camera_input("📱 调用摄像头拍照", key="atom_camera_img")
                    
                    a_content = st.text_area("📝 编辑内容 (支持 Markdown)", value=st.session_state.atom_content_prompt, height=150)
                    a_files = st.file_uploader("📎 上传附件", accept_multiple_files=True, key="atom_other_files")
                    submitted = st.form_submit_button("💾 保存笔记")
                    
                    if submitted and a_title:
                        file_paths = []
                        img_paths_str = ""
                        atom_dir = os.path.join(user_storage, "atom_notes_files")
                        
                        if upload_img:
                            raw_bytes = upload_img.getvalue()
                            if len(raw_bytes) > 2 * 1024 * 1024:
                                compressed = compress_image(raw_bytes)
                            else:
                                compressed = raw_bytes
                            save_path = os.path.join(atom_dir, upload_img.name)
                            with open(save_path, "wb") as f:
                                f.write(compressed)
                            file_paths.append(save_path)
                            img_paths_str += f"\n![{upload_img.name}]({save_path})\n"

                        if camera_img:
                            raw_bytes = camera_img.getvalue()
                            if len(raw_bytes) > 2 * 1024 * 1024:
                                compressed = compress_image(raw_bytes)
                            else:
                                compressed = raw_bytes
                            cam_fname = f"camera_{datetime.now().strftime('%Y%m%d%H%M%S')}.jpg"
                            save_path = os.path.join(atom_dir, cam_fname)
                            with open(save_path, "wb") as f:
                                f.write(compressed)
                            file_paths.append(save_path)
                            img_paths_str += f"\n![{cam_fname}]({save_path})\n"
                        
                        if a_files:
                            for file in a_files:
                                save_path = os.path.join(atom_dir, file.name)
                                with open(save_path, "wb") as f:
                                    f.write(file.getbuffer())
                                file_paths.append(save_path)
                        
                        file_paths_str = ",".join(file_paths)
                        final_content = f"{img_paths_str}\n\n{a_content}"
                        
                        conn = get_db_connection()
                        c = conn.cursor()
                        c.execute("INSERT INTO atom_notes (user_id, title, content, file_paths, code_content, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                                  (user['id'], a_title, final_content, file_paths_str, "", datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                        conn.commit()
                        conn.close()
                        st.success("✅ 笔记已保存！")
                        st.rerun()

        with tab_flow:
            st.subheader("⚡ 可视化白板流程图 (自由拖拽、连线、贴图)")
            st.caption("✅ 从左侧工具栏拖拽节点；双击编辑文字；拖拽连接点连线；右键节点查看更多操作。")

            col_add, col_save = st.columns([1, 1])
            with col_add:
                if st.button("➕ 添加节点 (自动编号)", use_container_width=True):
                    st.session_state.flow_json = "add_node"
            with col_save:
                if st.button("💾 保存当前流程图到笔记", use_container_width=True):
                    st.session_state.flow_json = "save_flow"

            drawflow_html = f"""
            <style>
                #drawflow-container {{
                    height: 600px;
                    background: #f7f9fc;
                    border: 1px solid #e0e0e0;
                    border-radius: 8px;
                    position: relative;
                    overflow: hidden;
                }}
                .drawflow-node {{
                    width: 150px !important;
                    background: #ffffff !important;
                    border: 2px solid #3b82f6 !important;
                    border-radius: 8px !important;
                    box-shadow: 0 4px 6px rgba(0,0,0,0.05) !important;
                    padding: 10px !important;
                }}
                .drawflow-node .node-content {{
                    text-align: center;
                    font-family: sans-serif;
                }}
                .drawflow-node img.node-img {{
                    max-width: 100%; 
                    height: auto; 
                    max-height: 80px; 
                    object-fit: contain; 
                    margin-top: 5px;
                    cursor: pointer;
                    border-radius: 4px;
                    border: 1px solid #eee;
                }}
                .drawflow-node .node-text {{
                    font-weight: bold;
                    color: #333;
                    margin-bottom: 5px;
                }}
                #toolbar {{
                    position: absolute;
                    top: 10px;
                    left: 10px;
                    z-index: 10;
                    display: flex;
                    gap: 8px;
                }}
                #toolbar button {{
                    background: #fff;
                    border: 1px solid #ccc;
                    border-radius: 4px;
                    padding: 6px 12px;
                    cursor: pointer;
                    font-size: 14px;
                }}
                #node-library {{
                    position: absolute;
                    top: 60px;
                    left: 10px;
                    z-index: 10;
                    background: rgba(255,255,255,0.9);
                    border: 1px solid #ccc;
                    border-radius: 8px;
                    padding: 8px;
                    display: flex;
                    flex-direction: column;
                    gap: 6px;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                }}
                #node-library .node-btn {{
                    width: 40px;
                    height: 40px;
                    border: 1px solid #ddd;
                    border-radius: 4px;
                    background: #fff;
                    cursor: pointer;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 20px;
                }}
                #mobile-fab {{
                    display: none;
                    position: absolute;
                    bottom: 20px;
                    right: 20px;
                    z-index: 20;
                    width: 56px;
                    height: 56px;
                    border-radius: 50%;
                    background: #3b82f6;
                    color: white;
                    font-size: 32px;
                    border: none;
                    box-shadow: 0 4px 10px rgba(0,0,0,0.2);
                    cursor: pointer;
                }}
                #mobile-node-menu {{
                    display: none;
                    position: absolute;
                    bottom: 90px;
                    right: 20px;
                    z-index: 20;
                    background: white;
                    border: 1px solid #ddd;
                    border-radius: 8px;
                    padding: 8px;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.15);
                    flex-direction: column;
                    gap: 6px;
                }}
                #mobile-node-menu .node-btn {{
                    width: 48px;
                    height: 48px;
                    border: none;
                    background: transparent;
                    font-size: 24px;
                    cursor: pointer;
                    border-radius: 4px;
                }}
                @media (max-width: 600px) {{
                    #node-library {{ display: none; }}
                    #mobile-fab {{ display: block; }}
                    #mobile-node-menu.active {{ display: flex; }}
                }}
                #loading {{
                    position: absolute;
                    top: 50%; left: 50%;
                    transform: translate(-50%, -50%);
                    font-size: 18px; color: #666;
                }}
            </style>
            <link rel="stylesheet" href="./static/drawflow.min.css">
            <script src="./static/drawflow.min.js"></script>
            <div id="drawflow-container">
                <div id="loading">加载中...</div>
                <div id="toolbar">
                    <button id="btn-clear" title="清空画布">🗑️</button>
                    <button id="btn-fit" title="适配视图">🔍</button>
                    <button id="btn-save" title="保存">💾</button>
                    <button id="btn-export" title="导出 PNG">📷</button>
                </div>
                <div id="node-library">
                    <button class="node-btn" data-type="rect" title="矩形">▭</button>
                    <button class="node-btn" data-type="rect-round" title="圆角矩形">▭</button>
                    <button class="node-btn" data-type="diamond" title="菱形">◇</button>
                    <button class="node-btn" data-type="circle" title="圆形">●</button>
                </div>
                <button id="mobile-fab">+</button>
                <div id="mobile-node-menu">
                    <button class="node-btn" data-type="rect" title="矩形">▭</button>
                    <button class="node-btn" data-type="rect-round" title="圆角矩形">▭</button>
                    <button class="node-btn" data-type="diamond" title="菱形">◇</button>
                    <button class="node-btn" data-type="circle" title="圆形">●</button>
                </div>
            </div>
            <script>
                (function() {{
                    const container = document.getElementById('drawflow-container');
                    const loading = document.getElementById('loading');
                    let editor = null;
                    function initDrawflow() {{
                        editor = new Drawflow(container);
                        editor.start();
                        editor.zoom_ = 1; editor.zoom_max = 3; editor.zoom_min = 0.1;
                        document.querySelector('.drawflow-delete')?.remove();
                        const jsonStr = {st.session_state.flow_json};
                        if (jsonStr && jsonStr !== '{{}}' && jsonStr !== 'add_node' && jsonStr !== 'save_flow') {{
                            try {{ editor.import(JSON.parse(jsonStr)); }} catch(e) {{ console.warn('加载失败', e); }}
                        }}
                        loading.style.display = 'none';
                        bindEvents();
                    }}
                    function bindEvents() {{
                        document.querySelectorAll('.node-btn').forEach(btn => {{
                            btn.addEventListener('click', function(e) {{ addNodeByType(this.dataset.type); }});
                        }});
                        const fab = document.getElementById('mobile-fab');
                        const menu = document.getElementById('mobile-node-menu');
                        fab.addEventListener('click', function() {{ menu.classList.toggle('active'); }});
                        document.getElementById('btn-clear').addEventListener('click', function() {{ if(confirm('确定清空？')) editor.clear(); }});
                        document.getElementById('btn-fit').addEventListener('click', function() {{ editor.zoom_to_fit(); }});
                        document.getElementById('btn-save').addEventListener('click', saveData);
                        document.getElementById('btn-export').addEventListener('click', function() {{ alert('导出功能待完善'); }});
                        container.addEventListener('dblclick', function(e) {{
                            const node = e.target.closest('.drawflow-node');
                            if (node) {{
                                const id = node.id.replace('node-', '');
                                const textEl = node.querySelector('.node-text');
                                if (textEl) {{
                                    const newText = prompt('修改节点文字：', textEl.textContent);
                                    if (newText !== null) {{ textEl.textContent = newText; saveData(); }}
                                }}
                            }}
                        }});
                        container.addEventListener('contextmenu', function(e) {{
                            const node = e.target.closest('.drawflow-node');
                            if (!node) return;
                            e.preventDefault();
                            const id = node.id.replace('node-', '');
                            const textEl = node.querySelector('.node-text');
                            showContextMenu(e.clientX, e.clientY, id, textEl);
                        }});
                        window.addEventListener('message', function(e) {{
                            const data = e.data;
                            if (data === 'add_node') addNodeByType('rect');
                            else if (data === 'save_flow') saveData();
                        }});
                        editor.on('connectionCreated', saveData);
                        editor.on('nodeMoved', saveData);
                        editor.on('nodeDeleted', saveData);
                    }}
                    function addNodeByType(type) {{
                        let html = '', className = '';
                        switch(type) {{
                            case 'rect': html = `<div class="node-content"><div class="node-text">矩形</div></div>`; break;
                            case 'rect-round': html = `<div class="node-content"><div class="node-text">圆角</div></div>`; className = 'drawflow-node-round'; break;
                            case 'diamond': html = `<div class="node-content"><div class="node-text">判断</div></div>`; className = 'drawflow-node-diamond'; break;
                            case 'circle': html = `<div class="node-content"><div class="node-text">数据</div></div>`; className = 'drawflow-node-circle'; break;
                            default: html = `<div class="node-content"><div class="node-text">节点</div></div>`;
                        }}
                        const x = 100 + Math.random() * 300, y = 100 + Math.random() * 200;
                        editor.addNode('custom', 1, 1, x, y, 'custom', {{ html: html }}, null, className);
                        saveData();
                    }}
                    let contextMenu = null;
                    function showContextMenu(x, y, nodeId, textEl) {{
                        if (contextMenu) contextMenu.remove();
                        const menu = document.createElement('div');
                        menu.style.cssText = `position:absolute;background:white;border:1px solid #ccc;border-radius:4px;box-shadow:0 2px 8px rgba(0,0,0,0.15);padding:6px 0;z-index:100;min-width:120px;left:${{x}}px;top:${{y}}px;`;
                        const items = [
                            {{ label: '重命名', action: () => {{ const t=prompt('重命名：',textEl.textContent); if(t!==null){{textEl.textContent=t;saveData();}} }} }},
                            {{ label: '上传图片', action: () => {{ const input=document.createElement('input'); input.type='file';input.accept='image/*'; input.onchange=function(e){{ const file=e.target.files[0];if(file){{ const r=new FileReader(); r.onload=function(ev){{ const img=document.createElement('img'); img.className='node-img'; img.src=ev.target.result; document.querySelector(`#node-${{nodeId}} .node-content`).appendChild(img); saveData(); }}; r.readAsDataURL(file); }} }}; input.click(); }} }},
                            {{ label: '复制节点', action: () => {{ const newNode = editor.exportNode(nodeId); editor.addNode('custom', 1, 1, 100, 100, 'custom', newNode); saveData(); }} }},
                            {{ label: '删除节点', action: () => {{ if(confirm('确定删除？')){{ editor.removeNodeId(nodeId); saveData(); }} }} }}
                        ];
                        items.forEach(item => {{
                            const div=document.createElement('div'); div.textContent=item.label; div.style.cssText='padding:6px 12px;cursor:pointer;';
                            div.addEventListener('mouseenter',()=>div.style.background='#f0f0f0');
                            div.addEventListener('mouseleave',()=>div.style.background='transparent');
                            div.addEventListener('click',()=>{{ item.action(); menu.remove(); }});
                            menu.appendChild(div);
                        }});
                        document.body.appendChild(menu); contextMenu=menu;
                        setTimeout(()=>{{ document.addEventListener('click',function close(e){{ if(!menu.contains(e.target)){{ menu.remove(); document.removeEventListener('click',close); }} }}); }},10);
                    }}
                    let saveTimer = null;
                    function saveData() {{
                        if (!editor) return;
                        clearTimeout(saveTimer);
                        saveTimer = setTimeout(() => {{
                            const json = JSON.stringify(editor.export());
                            window.parent.postMessage({{ action: 'save_data', json: json }}, '*');
                        }}, 300);
                    }}
                    initDrawflow();
                }})();
            </script>
            """
            st.components.v1.html(drawflow_html, height=650)
            flow_code = st.text_area("流程图代码 (JSON备份)", value=st.session_state.flow_json, key="flow_json_code", height=100, label_visibility="collapsed")
            if st.button("保存流程图至本地数据库", type="primary"):
                if flow_code and flow_code != "{}" and flow_code != "add_node" and flow_code != "save_flow":
                    conn = get_db_connection()
                    c = conn.cursor()
                    flow_title = st.text_input("为这个流程图命名", value=f"流程图_{datetime.now().strftime('%Y%m%d%H%M')}")
                    if flow_title:
                        c.execute("INSERT INTO atom_notes (user_id, title, content, file_paths, code_content, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                                  (user['id'], flow_title, "该笔记为流程图白板数据", "", flow_code, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
                        conn.commit()
                        conn.close()
                        st.success("✅ 流程图已存入您的原子笔记！可随时在「图文笔记」中点击查看详情调出。")
                        st.rerun()

    # ---------- 个人中心 ----------
    elif menu == "⚙️ 个人中心":
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("SELECT identity_code, created_at FROM users WHERE id = ?", (user['id'],))
        user_info = c.fetchone()
        conn.close()
        
        st.write(f"**用户名**：{user['username']}")
        st.write(f"**您的专属 18 位身份码**：`{user_info[0]}`")
        st.write(f"**注册时间**：{user_info[1]}")
        st.write("**云端说明**：若部署在 Streamlit Cloud，数据默认保存在临时硬盘。如需数据云端持久化，请联系技术团队修改数据库为 PostgreSQL 并配置 OSS 存储。")

# ================= 路由控制 =================
if "user" not in st.session_state:
    if "page" not in st.session_state:
        st.session_state["page"] = "login"
    if st.session_state["page"] == "login":
        login_page()
    else:
        register_page()
else:
    main_app()
